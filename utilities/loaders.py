import pandas as pd
import os
import json
import datetime as dt
import re

from dateutil import parser 
from concurrent.futures import ThreadPoolExecutor
from collections import Counter
from docx import Document
from striprtf.striprtf import rtf_to_text
from PIL import Image

import fitz
import io
import pytesseract

def extract_images_text(pdf):
    """
    works on any pdf file flattened or unflattened by
    extracting any text or image with text within the pdf file
    """
    # define path to tessearact executable
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    # iterate over PDF pages
    data = []
    for page_index, page in enumerate(pdf):

        # get images on the page
        image_list = page.get_images(full=True)

        # if page contains no text then this statement will return 
        # an empty list akin to what page.get_images() naturally returns
        text_list = [] if page.get_text() == "" else page.get_text().split('\n')
        # print(text_list)
        # print(image_list)

        # printing number of images found in this page
        if image_list:
            print(f"[+] Found a total of {len(image_list)} images on page {page_index}")
        else:
            print("[!] No images found on page", page_index)

        data.extend(image_list)
        data.extend(text_list)

    # filter data for any duplicate tuples or whole strings
    data_filt = list(set(data))

    output = []
    for data_index, data in enumerate(data_filt, start=1):

        # note that data may be in a form of a tuple meaning an
        # image or just a full string. If such is the case that it
        # is a tuple proceed with extracting image from pdf and bytes
        # object and read it through pillow then convert to text via 
        # tesseract
        # print(type(data))
        if type(data) == tuple:
            # get the XREF of the image
            xref = data[0]

            # extract the image bytes
            image_obj = pdf.extract_image(xref)
            image_bytes = image_obj["image"]

            # get the image extension
            # image_ext = image_obj["ext"]

            # convert the bytes of the image to BytesIO object
            # so it can be read by Image.open() function
            base_image = Image.open(io.BytesIO(image_bytes))

            # return value will naturally be a giant string with \n char
            # so split it according to \n char to reveal lines
            text = pytesseract.image_to_string(base_image)
            text = text.split('\n')
            output.extend(text)

        elif type(data) == str:
            text = data.split('\n')
            output.extend(text)
    
    return output



def read_files(input_dir, files):
    """
    inspects the input directory of all types of files and outputs
    a dictionary with all the files categorized/bucketed into their
    different file types

    args:
        input_dir - is the directory of the files
        files - a list containing the names of the files in the input
        directory

    output:
    {
        'jsons': {
            'dataset#234': {<jsonfile>},
            'dataset#234': {<jsonfile>},
            'dataset#234': {<jsonfile>},
            ...
            'dataset#234': {<jsonfile>},
        },
        'txts': {
            'dataset#2324': {<read txt files>},

        },
        'docs': {
            'dataset#2342': {<read rtf or docx file>},
        }
        'pdfs: {
            'dataset#1231': {<read pdf file>}
        }
    }
    """

    def create_buckets(file_name):
        extension = re.search(r'(.json|.txt|.rtf|.docx|.csv|.xlsx|.md|.pdf|.webp|.png|.jpg)$', file_name)
        extension = extension[0] if extension != None else ""
        return extension

    # create buckets for each file type
    with ThreadPoolExecutor() as exe:
        out = list(set(exe.map(create_buckets, files)))
    buckets = {bucket: list(filter(lambda file_name: file_name.endswith(bucket), files)) for bucket in out}

    # read files
    def helper(file_name, bucket_name):
        # define path
        path = f'{input_dir}/{file_name}'
        
        # if for example the bucket being processed is .json then we 
        # only allow files to be returned that are .json
        if file_name.endswith('.json') and file_name.endswith(bucket_name):
            with open(path, 'r') as file:
                data = json.load(file)
                file.close()

            new_name = re.sub(r'.json$', '', file_name)
            # print(bucket_name, new_name)
            return new_name, data

        elif file_name.endswith('.txt') and file_name.endswith(bucket_name):
            with open(path, 'r') as file:
                data = file.readlines()
                file.close()

            new_name = re.sub(r'.txt$', '', file_name)
            # print(bucket_name, new_name)
            return new_name, data

        elif file_name.endswith('.rtf') and file_name.endswith(bucket_name):
            with open(path, 'r') as file:
                rtf_data = file.read()
                data = rtf_to_text(rtf_data)
                data = data.split('\n')
                file.close()

            new_name = re.sub(r'.rtf$', '', file_name)
            # print(bucket_name, new_name)
            return new_name, data

        elif file_name.endswith('.docx') and file_name.endswith(bucket_name):
            doc = Document(path)
            data = [doc.text for doc in doc.paragraphs]

            new_name = re.sub(r'.docx$', '', file_name)
            # print(bucket_name, new_name)
            return new_name, data
        
        elif file_name.endswith('.csv') and file_name.endswith(bucket_name):
            # note there are exceptions in reading .csv files in that 
            # sometimes there are other extensions of files that do indeed need
            # to be loaded via pandas i.e.
            # ds_5265 = pd.read_csv(f'./{input_dir}/Dataset#5265expanded_overtime_markets_data.txt', delimiter='|')
            df = pd.read_csv(path)
            new_name = re.sub(r'.csv$', '', file_name)
            # print(bucket_name, new_name)
            return new_name, df
        
        elif file_name.endswith('.xlsx') and file_name.endswith(bucket_name):
            df = pd.read_excel(path)
            new_name = re.sub(r'.xlsx$', '', file_name)
            # print(bucket_name, new_name)
            return new_name, df
        
        elif file_name.endswith('.pdf') and file_name.endswith(bucket_name):
            pdf = fitz.open(path)
            new_name = re.sub(r'.pdf$', '', file_name)

            # extract text from images and text in read pdf file
            data = extract_images_text(pdf)
            
            return new_name, data
        
    # O(n) processing
    outputs = {}
    for bucket_name, bucket_files in buckets.items():
        print(bucket_name, bucket_files)
        # process the files in under a specific bucket
        with ThreadPoolExecutor() as exe:

            # if for example .csv is the current bucket being processed
            # then the files that will only be returned for this are the .csvs
            n_files_in_bucket = len(bucket_files)
            outputs[bucket_name] = list(exe.map(
                helper, 
                bucket_files, 
                [bucket_name] * n_files_in_bucket
            ))


    return outputs



def read_doc_files(input_dir: str, files: list[str]) -> dict[list[dict]]:
    def helper(file_name):
        path = f'{input_dir}/{file_name}'
        doc = Document(path)
        data = [doc.text for doc in doc.paragraphs]

        new_name = re.sub(r'.docx|.rtf$', '', file_name)
        return new_name, data

    # concurrently read and load all .json files
    with ThreadPoolExecutor() as exe:
        docs = dict(list(exe.map(helper, files)))
    return docs



def read_json_files(input_dir: str, files: list[str]) -> dict[list[dict]]:
    def helper(file_name):
        with open(f'{input_dir}/{file_name}', 'r') as file:
            data = json.load(file)
            file.close()

        new_name = file_name.replace('.json', '')
        return new_name, data

    # concurrently read and load all .json files
    with ThreadPoolExecutor() as exe:
        jsons = dict(list(exe.map(helper, files)))
    return jsons



def read_txt_files(input_dir: str, files: list[str]) -> dict[list[str]]: 
    def helper(file_name):
        with open(f'{input_dir}/{file_name}', 'r') as file:
            data = file.readlines()
            file.close()

        new_name = re.sub(r'.txt|.md$', '', file_name)
        return new_name, data

    # concurrently read and load all .json files
    with ThreadPoolExecutor() as exe:
        txts = dict(list(exe.map(helper, files)))
    return txts



def read_rtf_files(input_dir: str, files: list[str]) -> dict[list[str]]:
    def helper(file_name):
        with open(f'{input_dir}/{file_name}', 'r') as file:
            rtf_data = file.read()
            data = rtf_to_text(rtf_data)
            data = data.split('\n')
            file.close()

        new_name = file_name.replace('.rtf', '')
        return new_name, data

    # concurrently read and load all .json files
    with ThreadPoolExecutor() as exe:
        rtfs = dict(list(exe.map(helper, files)))
    return rtfs

def read_xlsx_files(input_dir: str, files: list[str]) -> dict[list[str]]:
    def helper(file_name):
        df = pd.read_excel(f'{input_dir}/{file_name}')
        new_name = re.sub(r'.xlsx$', '', file_name)

        return new_name, df
    
    # concurrently read and load all .xlsx files
    with ThreadPoolExecutor() as exe:
        xlsxs = dict(list(exe.map(helper, files)))
    return xlsxs

def read_csv_files(input_dir: str, files: list[str]) -> dict[list[str]]:
    def helper(file_name):
        df = pd.read_csv(f'{input_dir}/{file_name}')
        new_name = re.sub(r'.csv$', '', file_name)

        return new_name, df
    
    # concurrently read and load all .xlsx files
    with ThreadPoolExecutor() as exe:
        csvs = dict(list(exe.map(helper, files)))
    return csvs